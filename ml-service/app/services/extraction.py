"""Document text extraction service.

Supports PDF, DOCX, XLSX, CSV, and images (via OCR).
"""

import io
import fitz  # PyMuPDF
import pdfplumber
from docx import Document as DocxDocument
import pandas as pd
import pytesseract
from PIL import Image


class ExtractionResult:
    def __init__(self, text: str, tables: list[list[list[str]]], images: list[bytes], metadata: dict):
        self.text = text
        self.tables = tables
        self.images = images
        self.metadata = metadata


def extract_pdf(file_bytes: bytes) -> ExtractionResult:
    """Extract text, tables, and images from a PDF."""
    # Text extraction via PyMuPDF (fast, layout-aware)
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text_parts = []
    images = []
    metadata = {
        "page_count": len(doc),
        "title": doc.metadata.get("title", ""),
        "author": doc.metadata.get("author", ""),
    }

    for page_num, page in enumerate(doc):
        page_text = page.get_text("text")
        if page_text.strip():
            text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
        else:
            # Fallback to OCR for scanned pages
            pix = page.get_pixmap(dpi=200)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            ocr_text = pytesseract.image_to_string(img)
            if ocr_text.strip():
                text_parts.append(f"[Page {page_num + 1} (OCR)]\n{ocr_text}")

        # Extract images
        for img_info in page.get_images(full=True):
            xref = img_info[0]
            base_image = doc.extract_image(xref)
            if base_image:
                images.append(base_image["image"])

    doc.close()

    # Table extraction via pdfplumber (better for structured tables)
    tables = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_tables = page.extract_tables()
            for table in page_tables:
                if table:
                    tables.append(table)

    return ExtractionResult(
        text="\n\n".join(text_parts),
        tables=tables,
        images=images,
        metadata=metadata,
    )


def extract_docx(file_bytes: bytes) -> ExtractionResult:
    """Extract text from a DOCX file."""
    doc = DocxDocument(io.BytesIO(file_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            prefix = ""
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "")
                prefix = f"{'#' * int(level)} " if level.isdigit() else "## "
            paragraphs.append(f"{prefix}{para.text}")

    # Extract tables
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        tables.append(rows)

    return ExtractionResult(
        text="\n\n".join(paragraphs),
        tables=tables,
        images=[],
        metadata={"paragraph_count": len(paragraphs)},
    )


def extract_spreadsheet(file_bytes: bytes, filename: str) -> ExtractionResult:
    """Extract data from XLSX or CSV files."""
    if filename.endswith(".csv"):
        df = pd.read_csv(io.BytesIO(file_bytes))
    else:
        df = pd.read_excel(io.BytesIO(file_bytes))

    # Convert to readable text representation
    text_parts = [
        f"Spreadsheet: {filename}",
        f"Shape: {df.shape[0]} rows x {df.shape[1]} columns",
        f"Columns: {', '.join(df.columns.tolist())}",
        "",
        "Data Summary:",
        str(df.describe()),
        "",
        "First 20 rows:",
        df.head(20).to_string(),
    ]

    return ExtractionResult(
        text="\n".join(text_parts),
        tables=[df.values.tolist()],
        images=[],
        metadata={
            "row_count": df.shape[0],
            "col_count": df.shape[1],
            "columns": df.columns.tolist(),
        },
    )


def extract_image(file_bytes: bytes) -> ExtractionResult:
    """Extract text from an image via OCR."""
    img = Image.open(io.BytesIO(file_bytes))
    ocr_text = pytesseract.image_to_string(img)

    return ExtractionResult(
        text=ocr_text,
        tables=[],
        images=[file_bytes],
        metadata={"width": img.width, "height": img.height, "mode": img.mode},
    )


def extract_document(file_bytes: bytes, mime_type: str, filename: str = "") -> ExtractionResult:
    """Route to the appropriate extractor based on MIME type."""
    if mime_type == "application/pdf":
        return extract_pdf(file_bytes)
    elif mime_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        return extract_docx(file_bytes)
    elif mime_type in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
        "text/csv",
    ):
        return extract_spreadsheet(file_bytes, filename)
    elif mime_type.startswith("image/"):
        return extract_image(file_bytes)
    else:
        # Fallback: treat as plain text
        text = file_bytes.decode("utf-8", errors="replace")
        return ExtractionResult(text=text, tables=[], images=[], metadata={})
